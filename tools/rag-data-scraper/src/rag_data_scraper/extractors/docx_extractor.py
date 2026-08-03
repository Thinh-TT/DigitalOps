import hashlib
from pathlib import Path
from typing import List
import docx
from .base import BaseExtractor, ExtractedDocument, ContentBlock, BlockType

class DOCXExtractor(BaseExtractor):
    def extract(self, file_path: Path | str) -> ExtractedDocument:
        path = Path(file_path)
        with open(path, "rb") as f:
            raw_bytes = f.read()
        raw_sha256 = hashlib.sha256(raw_bytes).hexdigest()

        doc = docx.Document(path)
        blocks: List[ContentBlock] = []
        title = path.stem

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = p.style.name.lower() if p.style else ""
            if "heading 1" in style_name or "heading 2" in style_name or "heading 3" in style_name:
                level = 1
                if "heading 2" in style_name:
                    level = 2
                elif "heading 3" in style_name:
                    level = 3
                blocks.append(ContentBlock(
                    block_type=BlockType.HEADING,
                    text=text,
                    heading_level=level
                ))
            else:
                blocks.append(ContentBlock(
                    block_type=BlockType.PARAGRAPH,
                    text=text
                ))

        for table in doc.tables:
            table_grid = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_grid.append(row_cells)
            if table_grid:
                table_text = "\n".join([" | ".join(r) for r in table_grid])
                blocks.append(ContentBlock(
                    block_type=BlockType.TABLE,
                    text=table_text,
                    table_data=table_grid
                ))

        return ExtractedDocument(
            source_uri=str(path.resolve()),
            title=title,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            raw_sha256=raw_sha256,
            blocks=blocks,
            ocr_used=False,
            ocr_confidence=1.0
        )
